import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


def set_cell_text(cell, text):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text.strip())
    run.font.name = "SimSun"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    run.font.size = Pt(10.5)


def set_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), "000000")
        borders.append(tag)
    tbl_pr.append(borders)


def add_paragraph(document, text, style=None):
    paragraph = document.add_paragraph(style=style)
    paragraph.paragraph_format.first_line_indent = Cm(0.74) if not style else None
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    run.font.name = "SimSun"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    run.font.size = Pt(12)
    return paragraph


def add_code_block(document, lines):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(0.7)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run("\n".join(lines))
    run.font.name = "Courier New"
    run.font.size = Pt(9)


def apply_styles(document):
    normal = document.styles["Normal"]
    normal.font.name = "SimSun"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    normal.font.size = Pt(12)

    for style_name in ("Heading 1", "Heading 2", "Heading 3", "Heading 4"):
        style = document.styles[style_name]
        style.font.name = "SimHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "SimHei")
        style.font.bold = True


def parse_inline(text):
    return re.sub(r"`([^`]+)`", r"\1", text).strip()


def convert(md_path, docx_path):
    text = Path(md_path).read_text(encoding="utf-8")
    lines = text.splitlines()

    document = Document()
    apply_styles(document)
    section = document.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

    i = 0
    in_code = False
    code_lines = []
    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()

        if line.startswith("```"):
            if in_code:
                add_code_block(document, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(raw)
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        if line.startswith("|") and i + 1 < len(lines) and lines[i + 1].startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i])
                i += 1
            rows = []
            for table_line in table_lines:
                cells = [parse_inline(c) for c in table_line.strip().strip("|").split("|")]
                if all(re.fullmatch(r"\s*:?-{3,}:?\s*", c or "") for c in cells):
                    continue
                rows.append(cells)
            if rows:
                table = document.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = "Table Grid"
                set_table_borders(table)
                for r, row in enumerate(rows):
                    for c, value in enumerate(row):
                        set_cell_text(table.cell(r, c), value)
                        if r == 0:
                            for run in table.cell(r, c).paragraphs[0].runs:
                                run.bold = True
                document.add_paragraph()
            continue

        heading = re.match(r"^(#{1,4})\s+(.+)$", line)
        if heading:
            level = len(heading.group(1))
            title = parse_inline(heading.group(2))
            paragraph = document.add_heading(title, level=level)
            for run in paragraph.runs:
                run.font.name = "SimHei"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimHei")
            if level == 1:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        ordered = re.match(r"^\d+\.\s+(.+)$", line)
        if ordered:
            paragraph = document.add_paragraph(style="List Number")
            paragraph.paragraph_format.line_spacing = 1.5
            run = paragraph.add_run(parse_inline(ordered.group(1)))
            run.font.name = "SimSun"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
            run.font.size = Pt(12)
            i += 1
            continue

        add_paragraph(document, parse_inline(line))
        i += 1

    document.save(docx_path)


if __name__ == "__main__":
    if len(sys.argv) != 3:
        raise SystemExit("usage: md_to_docx.py input.md output.docx")
    convert(sys.argv[1], sys.argv[2])
